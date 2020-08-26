from docx import Document


def convert_amino_acid_abbreviation(three_letter_abbreviation):

    one_letter_abbreviation = {
        'ala': 'A',
        'arg': 'R',
        'asn': 'N',
        'asp': 'D',
        'cys': 'C',
        'gln': 'Q',
        'glu': 'E',
        'gly': 'G',
        'his': 'H',
        'ile': 'I',
        'leu': 'L',
        'lys': 'K',
        'met': 'M',
        'phe': 'F',
        'pro': 'P',
        'ser': 'S',
        'thr': 'T',
        'trp': 'W',
        'tyr': 'Y',
        'val': 'V',
        'asx': 'B',
        'glx': 'Z',
        'xaa': 'X',  # unknow or other
    }.get(three_letter_abbreviation.lower())

    if not one_letter_abbreviation:
        raise Exception(f'Cannot find One-letter abbreviation of {three_letter_abbreviation}')

    return one_letter_abbreviation


def convert_sequence(sequence_listing_file):

    SEQ_INDEX_LABEL = '<210>'
    SEQ_LEN_LABEL = '<211>'
    SEQ_TYPE_LABEL = '<212>'
    SEQ_CONTENT_LABEL = '<400>'
    PRT_TYPE = 'PRT'

    seq_dict = {}
    while True:
        line = sequence_listing_file.readline().decode()
        if not line:
            break

        if SEQ_INDEX_LABEL in line:
            seq_index = int(line.split(SEQ_INDEX_LABEL)[1].strip())
            try:
                seq_len = int(sequence_listing_file.readline().decode().split(SEQ_LEN_LABEL)[1].strip())
                seq_type = sequence_listing_file.readline().decode().split(SEQ_TYPE_LABEL)[1].strip()
            except IndexError:
                # 210, 211, 212 must appear in three consecutive lines
                continue

            dna_or_rna_seq = prt_seq = ''
            while True:
                line = sequence_listing_file.readline().decode()
                if SEQ_CONTENT_LABEL in line:
                    target_seq = ''
                    while len(target_seq) < seq_len:
                        line = sequence_listing_file.readline().decode()
                        for fragment in line.split(' '):
                            if fragment.isalpha():
                                if fragment.islower():
                                    dna_or_rna_seq += fragment.strip()
                                else:
                                    prt_seq += convert_amino_acid_abbreviation(fragment)
                        target_seq = prt_seq if seq_type.upper() == PRT_TYPE else dna_or_rna_seq
                    break

            seq_dict.update({
                seq_index: {
                    'len': seq_len,
                    'type': seq_type,
                    'dna_or_rna_seq': dna_or_rna_seq,
                    'prt_seq': prt_seq,
                    'is_found': False,
                },
            })

    return seq_dict


def compare_sequence(seq_dict, manual_file):

    document = Document(manual_file)

    for _, seq_content in seq_dict.items():
        if seq_content['is_found']:
            continue

        target_seqs = []
        dna_or_rna_seq = seq_content.get('dna_or_rna_seq')
        prt_seq = seq_content.get('prt_seq')
        if dna_or_rna_seq:
            target_seqs.append(dna_or_rna_seq)
        if prt_seq:
            target_seqs.append(prt_seq)

        seq_found = [False] * len(target_seqs)
        for index, target_seq in enumerate(target_seqs):
            for paragraph in document.paragraphs:
                if target_seq in paragraph.text:
                    seq_found[index] = True
                    break

        for index, target_seq in enumerate(target_seqs):
            if seq_found[index]:
                continue

            for table in document.tables:
                for cell in table._cells:
                    for paragraph in cell.paragraphs:
                        if target_seq in paragraph.text:
                            seq_found[index] = True
                            break
                    else:
                        continue
                    break
                else:
                    continue
                break

        seq_content['is_found'] = sum(seq_found) == len(target_seqs)
